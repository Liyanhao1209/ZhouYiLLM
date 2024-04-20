import json
from docx import Document

# 读取JSON文件
with open('self_qa.json', 'r', encoding='utf-8') as file:
    data = json.load(file)

# 创建Word文档并添加内容
doc = Document()
for item in data:
    question = item.get('instruction', '')
    answer = item.get('answer', '')
    doc.add_heading('问题:', level=2)
    doc.add_paragraph(question)
    doc.add_heading('回答:', level=2)
    doc.add_paragraph(answer)

# 保存Word文件
doc.save('output.docx')
